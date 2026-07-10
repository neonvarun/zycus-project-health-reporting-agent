from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "rag_methodology.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    properties = tbl.tblPr

    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_value))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DCE3", size="6"):
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def mark_header_row(row):
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_paragraph(document, text="", bold_prefix=None, style=None, color=None, italic=False):
    paragraph = document.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=10)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(document, label, text, fill=CALLOUT_FILL, label_color=INK):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D7DCE3")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(label + " ")
    set_run_font(label_run, size=10.5, color=label_color, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_document():
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("ZYCUS PROFESSIONAL SERVICES PMO")
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Project Health Reporting Agent  |  Confidential  |  Page ")
    set_run_font(footer_run, size=8.5, color=MUTED)
    set_page_field(footer)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("RAG METHODOLOGY FRAMEWORK")
    set_run_font(title_run, size=23, color="000000", bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Project Health Reporting Agent")
    set_run_font(subtitle_run, size=14, color="373737")

    metadata = add_table(
        document,
        ["Purpose", "Operating model", "Updated"],
        [[
            "Automated delivery health reporting",
            "Deterministic RAG + local AI explanation",
            date.today().strftime("%d %b %Y"),
        ]],
        [3200, 3900, 2260],
    )
    for cell in metadata.rows[1].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_callout(
        document,
        "Decision principle:",
        "Python owns the health score and final RAG status. The local Qwen model only turns evaluator-owned metrics into readable executive reasoning.",
    )

    document.add_heading("1. Objectives and design principles", level=1)
    add_bullet(document, "Deterministic core: the same project data produces the same RAG score and status.")
    add_bullet(document, "AI-enhanced reasoning: local Qwen3.5-0.8B Q4_K_M polishes explanations; Gemini remains an optional explicit provider.")
    add_bullet(document, "Data resilience: missing sheets, incomplete cells, formula errors, and unparseable dates are recorded as assumptions rather than causing the pipeline to crash.")

    document.add_heading("2. Weighted health score", level=1)
    add_paragraph(document, "The project health score is calculated on a 0–100 scale:")
    formula = add_callout(document, "Health score:", "0.30S + 0.25M + 0.20B + 0.10C + 0.10E + 0.05D", fill="E8EEF5", label_color=BLUE)
    add_table(
        document,
        ["Signal", "Weight", "What it measures"],
        [
            ("Schedule variance (S)", "30%", "Elapsed timeline versus physical completion."),
            ("Milestone health (M)", "25%", "Overdue milestones and their share of milestones."),
            ("Risks and blockers (B)", "20%", "On-hold, at-risk, red-health, and negatively commented tasks."),
            ("Budget proxy (C)", "10%", "SPI: completion divided by elapsed timeline."),
            ("Stakeholder sentiment (E)", "10%", "Negative keywords in linked or inline comments."),
            ("Data quality (D)", "5%", "Missing sheets, columns, and unparseable values."),
        ],
        [2800, 1100, 5460],
    )

    document.add_heading("3. Signal rules", level=1)
    document.add_heading("A. Schedule variance", level=2)
    add_paragraph(document, "Time elapsed % = (Today − Start Date) / (End Date − Start Date). Schedule slippage = Time elapsed % − Overall % Complete.")
    add_table(document, ["Score", "Rule"], [("100", "Slippage ≤ 5%."), ("50", "Slippage > 5% and ≤ 15%."), ("0", "Slippage > 15%.")], [1500, 7860])

    document.add_heading("B. Milestone health", level=2)
    add_paragraph(document, "An overdue milestone has an end date before today and is not marked Completed. The evaluator considers level 1–2 tasks and explicitly flagged milestones.")
    add_table(document, ["Score", "Rule"], [("100", "0 overdue milestones."), ("80", "1–2 overdue milestones or overdue ratio ≤ 5%."), ("50", "3–5 overdue milestones or overdue ratio ≤ 15%."), ("30", "More than 5 overdue milestones and ratio > 15%.")], [1500, 7860])

    document.add_heading("C. Risks and blockers", level=2)
    add_paragraph(document, "A task is treated as blocked when it is On Hold, marked At Risk, has Red schedule health, or contains a negative-risk comment.")
    add_table(document, ["Score", "Rule"], [("100", "0 active blockers."), ("80", "Up to 3 blockers or blocker ratio ≤ 3%."), ("50", "Up to 10 blockers or blocker ratio ≤ 6%."), ("30", "More than 10 blockers and ratio > 6%.")], [1500, 7860])

    document.add_heading("D. Budget proxy", level=2)
    add_paragraph(document, "Because standard plans do not contain a financial ledger, SPI is used as a schedule-based budget proxy: SPI = Overall % Complete / Time Elapsed %. If SPI cannot be computed, the score is neutral at 100.")
    add_table(document, ["Score", "Rule"], [("100", "SPI ≥ 0.95."), ("50", "SPI ≥ 0.80 and < 0.95."), ("0", "SPI < 0.80.")], [1500, 7860])

    document.add_heading("E. Stakeholder sentiment", level=2)
    add_paragraph(document, "Negative keywords include delay, delayed, impacted, pending, blocker, issue, risk, dependency, waiting, need, missing, failed, problem, concern, on hold, escalate, and stuck.")
    add_table(document, ["Score", "Rule"], [("100", "No negative comments or no comments."), ("80", "1–2 negative comments."), ("65", "3–5 negative comments."), ("50", "6–10 negative comments."), ("30", "More than 10 negative comments.")], [1500, 7860])

    document.add_heading("F. Data quality", level=2)
    add_paragraph(document, "The score starts at 100 and applies deductions for missing Summary sheets (−20), missing Comments sheets (−10), unparseable date values (−5 each), and missing core plan columns (−15). The final data-quality score cannot be below 0.")

    document.add_heading("4. Final RAG mapping and governance overrides", level=1)
    add_table(document, ["Overall score", "RAG", "Action"], [("80–100", "Green", "Maintain normal tracking."), ("60–79", "Amber", "Monitor risks and require mitigation."), ("Below 60", "Red", "Escalate for intervention and recovery planning.")], [1700, 1200, 6460])
    add_paragraph(document, "Overrides are applied after the weighted score:")
    add_number(document, "Red or Critical schedule health in the Summary sheet forces Red. Yellow or Amber forces Green results to Amber.")
    add_number(document, "Slippage above 25% forces Red.")
    add_number(document, "At least 15 blockers, blocker ratio at least 10%, and positive slippage above 5% forces Red.")

    document.add_heading("5. AI reasoning and resilience", level=1)
    add_bullet(document, "The default local provider downloads Qwen3.5-0.8B Q4_K_M on first use into .models/ and reuses the cached file on later runs.")
    add_bullet(document, "The model receives only evaluator-owned metrics and is instructed not to change the RAG status or invent facts.")
    add_bullet(document, "Generated output must contain Executive Summary, Why This Status Was Assigned, Top Risk Drivers, Recommended Actions, and Data Quality & Assumptions sections.")
    add_bullet(document, "If the model is unavailable, cannot download, or returns incomplete output, the rules-based explanation is returned and the report continues.")

    document.add_heading("6. Assumptions", level=1)
    add_bullet(document, "Project completion percentages are treated as physical progress percentages.")
    add_bullet(document, "Budget burn is inferred from schedule performance unless a future financial-ledger integration is added.")
    add_bullet(document, "Today’s date comes from the workbook when available; otherwise the runtime date is used.")
    add_bullet(document, "The local model requires internet access only for its first download and requires a compatible Python/native runtime.")

    document.core_properties.title = "RAG Methodology Framework"
    document.core_properties.subject = "Project Health Reporting Agent"
    document.core_properties.author = "Zycus Project Health Reporting Agent"
    document.core_properties.comments = "Generated from the repository methodology and executable scoring rules."
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
